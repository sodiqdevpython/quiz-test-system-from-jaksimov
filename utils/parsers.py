import docx

DOCX_AVAILABLE = True

class WordTestReader:
    def __init__(self):
        self.questions = []
        self.document = None

    def read_test_file(self, file_path):
        if not DOCX_AVAILABLE:
            return {'success': False, 'error': 'python-docx kutubxonasi o\'rnatilmagan', 'questions': []}

        try:
            self.document = docx.Document(file_path)
            print(
                f"Document ochildi. Paragraflar: {len(self.document.paragraphs)}, Jadvallar: {len(self.document.tables)}")

            questions = self._parse_document()
            return {'success': True, 'questions': questions, 'total_questions': len(questions), 'file_path': file_path}

        except Exception as e:
            print(f"Word fayl ochishda xatolik: {str(e)}")
            return {'success': False, 'error': str(e), 'questions': []}

    def _parse_document(self):
        questions = []

        if self.document.tables:
            print("Jadvallar topildi...")
            for table_idx, table in enumerate(self.document.tables):
                table_questions = self._parse_table_with_precise_images(table, table_idx)
                questions.extend(table_questions)

        print(f"Jami topilgan savollar: {len(questions)}")
        return questions

    def _parse_table_with_precise_images(self, table, table_idx):
        questions = []

        try:
            print(
                f"Jadval {table_idx} - {len(table.rows)} qator x {len(table.rows[0].cells) if table.rows else 0} ustun")

            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:
                    continue

                print(f"\nQator {row_idx} tahlil qilinmoqda...")
                cells = row.cells

                if len(cells) >= 5:
                    # Har bitta katakdan rasm bilan text ni ajratib oldin
                    question_cell = cells[1]
                    option_a_cell = cells[2]
                    option_b_cell = cells[3]
                    option_c_cell = cells[4]
                    option_d_cell = cells[5] if len(cells) > 5 else None

                    # Savolga matn bilan rasmnni olganim
                    question_text, question_image = self._extract_text_and_image(question_cell, "Savol")
                    if not question_text:
                        continue

                    # A variant
                    option_a_text, option_a_image = self._extract_text_and_image(option_a_cell, "A")
                    # B variant
                    option_b_text, option_b_image = self._extract_text_and_image(option_b_cell, "B")
                    # C variant
                    option_c_text, option_c_image = self._extract_text_and_image(option_c_cell, "C")
                    # D variant
                    option_d_text, option_d_image = None, None
                    if option_d_cell:
                        option_d_text, option_d_image = self._extract_text_and_image(option_d_cell, "D")

                    question = {
                        'id': cells[0].text.strip(),
                        'text': question_text,
                        'image_data': question_image,
                        'options': [
                            {'text': option_a_text or '', 'is_correct': True, 'image_data': option_a_image},
                            {'text': option_b_text or '', 'is_correct': False, 'image_data': option_b_image},
                            {'text': option_c_text or '', 'is_correct': False, 'image_data': option_c_image},
                            {'text': option_d_text or '', 'is_correct': False, 'image_data': option_d_image}
                        ]
                    }
                    questions.append(question)

                    # Debug agar xato chiqsa o'qishimga o'zimga
                    print(f"Savol yaratildi: '{question_text[:50]}...'")
                    if question_image:
                        print(f"Savolda rasm: {question_image['extension']} ({len(question_image['data'])} bayt)")

                    for i, opt in enumerate(question['options']):
                        opt_letter = chr(65 + i)  # A, B, C, D
                        print(f"    {opt_letter}) '{opt['text'][:30]}...' {'✓' if opt['is_correct'] else '✗'}")
                        if opt['image_data']:
                            print(
                                f"Rasm: {opt['image_data']['extension']} ({len(opt['image_data']['data'])} bayt)")

        except Exception as e:
            print(f"Jadval parse qilishda xato: {e}")
            import traceback
            traceback.print_exc()

        return questions

    def _extract_text_and_image(self, cell, cell_name):
        text = cell.text.strip()
        image = None

        try:
            print(f"{cell_name} katakchasi: '{text[:30]}...'")

            for para_idx, paragraph in enumerate(cell.paragraphs):
                for run_idx, run in enumerate(paragraph.runs):
                    if hasattr(run, '_element'):
                        drawings = run._element.xpath('.//w:drawing')
                        for drawing in drawings:
                            blips = drawing.xpath('.//a:blip', namespaces={
                                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                            })

                            for blip in blips:
                                rId = blip.get(
                                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if rId and rId in self.document.part.related_parts:
                                    image_part = self.document.part.related_parts[rId]
                                    if hasattr(image_part, 'blob'):
                                        image_data = image_part.blob

                                        if image_data.startswith(b'\xff\xd8'):
                                            ext = 'jpg'
                                        elif image_data.startswith(b'\x89PNG'):
                                            ext = 'png'
                                        elif image_data.startswith(b'GIF'):
                                            ext = 'gif'
                                        elif image_data.startswith(b'RIFF') and b'WEBP' in image_data[8:16]:
                                            ext = 'webp' #! buni yangi qo'shdim hali testlab ko'rmadim lekin ishlashi kerak
                                        else:
                                            ext = 'jpg'

                                        image = {
                                            'data': image_data,
                                            'extension': ext,
                                            'content_type': f'image/{ext}',
                                            'rId': rId
                                        }

                                        print(f"RASM TOPILDI: {ext} format, {len(image_data)} bayt, rId: {rId}")
                                        return text, image

            if not image:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if hasattr(run, '_element'):
                            pics = run._element.xpath('.//pic:pic', namespaces={
                                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
                            })

                            for pic in pics:
                                blips = pic.xpath('.//a:blip', namespaces={
                                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                                })

                                for blip in blips:
                                    rId = blip.get(
                                        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    if rId and rId in self.document.part.related_parts:
                                        image_part = self.document.part.related_parts[rId]
                                        if hasattr(image_part, 'blob'):
                                            image_data = image_part.blob

                                            if image_data.startswith(b'\xff\xd8'):
                                                ext = 'jpg'
                                            elif image_data.startswith(b'\x89PNG'):
                                                ext = 'png'
                                            elif image_data.startswith(b'RIFF') and b'WEBP' in image_data[8:16]:
                                                ext = 'webp' #! buni yangi qo'shdim hali testlab ko'rmadim lekin ishlashi kerak
                                            else:
                                                ext = 'jpg'

                                            image = {
                                                'data': image_data,
                                                'extension': ext,
                                                'content_type': f'image/{ext}',
                                                'rId': rId
                                            }

                                            print(f" PIC RASM: {ext}, {len(image_data)} bayt, rId: {rId}")
                                            return text, image

            if not image:
                print(f"Rasm topilmadi")

            return text, image

        except Exception as e:
            print(f"{cell_name} katakda xatolik: {e}")
            return text, None

    def debug_all_images(self):
        print("\n=== BARCHA RASMLARGA DEBUG ===")
        for part_id, part in self.document.part.related_parts.items():
            if hasattr(part, 'blob'):
                blob = part.blob
                if blob.startswith(b'\xff\xd8') or blob.startswith(b'\x89PNG') or blob.startswith(b'GIF'):
                    print(f"Rasm: {part_id} - {len(blob)} bayt")


def save_image_to_django(image_data, extension, prefix="question"):
    """Rasm ma'lumotlarini Django FileField uchun tayyorlash"""
    if not image_data:
        return None

    try:
        from django.core.files.base import ContentFile
        import uuid

        filename = f"{prefix}_{uuid.uuid4().hex[:8]}.{extension}"

        if isinstance(image_data, dict):
            raw_data = image_data.get('data')
        else:
            raw_data = image_data

        django_file = ContentFile(raw_data, name=filename)
        return django_file

    except Exception as e:
        print(f"Rasm saqlashda xato: {e}")
        return None


def parse_word_file_advanced(file_path):
    """Django signals uchun parser funksiya"""
    print(f"=== WORD FAYL PARSE QILISH BOSHLANDI ===")
    print(f"File path: {file_path}")

    try:
        if not DOCX_AVAILABLE:
            print("XATOLIK: python-docx kutubxonasi o'rnatilmagan!")
            return []

        reader = WordTestReader()
        result = reader.read_test_file(file_path)

        if result['success']:
            questions_data = []

            for q in result['questions']:
                question_data = {
                    'text': q['text'],
                    'image_data': q.get('image_data'),
                    'options': q['options']
                }
                questions_data.append(question_data)

            print(f"\nMuvaffaqiyatli parse qilindi: {len(questions_data)} ta savol")

            for i, q_data in enumerate(questions_data[:2]):
                print(f"\nSavol {i + 1}: {q_data['text'][:70]}...")
                if q_data.get('image_data'):
                    print(f"  Savolda rasm: {q_data['image_data']['extension']}")

                for j, opt in enumerate(q_data['options']):
                    has_text = "matn" if opt['text'] else ""
                    has_image = "rasm" if opt.get('image_data') else ""
                    content = f"{has_text}+{has_image}" if has_text and has_image else has_text or has_image or "bo'sh"
                    print(f"  {chr(65 + j)}) {content} ({'✓' if opt['is_correct'] else '✗'})")

            return questions_data
        else:
            print(f"Parse qilishda xatolik: {result['error']}")
            return []

    except Exception as e:
        print(f"Parser da umumiy xatolik: {str(e)}")
        import traceback
        traceback.print_exc()
        return []


def debug_word_images(file_path):
    try:
        reader = WordTestReader()
        reader.document = docx.Document(file_path)
        reader.debug_all_images()
    except Exception as e:
        print(f"Debug xatolik: {e}")


def debug_word_file(file_path):
    debug_word_images(file_path)

def parse_word_file_simple(file_path):
    return parse_word_file_advanced(file_path)

def test_parsing(file_path):
    return parse_word_file_advanced(file_path)