# test.py => bu aniq ishlaydigani

import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
import base64
import os

class WordTestReader:
    def __init__(self):
        self.questions = []

    def read_test_file(self, file_path):
        """
        Word faylini o'qish va test savollarini ajratib olish

        Args:
            file_path (str): Word fayl manzili

        Returns:
            dict: Test savollari va variantlari
        """
        try:
            # Word faylini ochish
            doc = docx.Document(file_path)

            # Faylni analyze qilish
            questions = self._parse_document(doc)

            return {
                'success': True,
                'questions': questions,
                'total_questions': len(questions),
                'file_path': file_path
            }

        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'questions': []
            }

    def _parse_document(self, doc):
        """Hujjatni parse qilish"""
        questions = []

        # Hujjatdagi barcha elementlarni tekshirish
        for element in doc.element.body:
            if isinstance(element, CT_Tbl):
                # Jadval formatidagi savollar
                table = Table(element, doc)
                table_questions = self._parse_table(table)
                questions.extend(table_questions)
            elif isinstance(element, CT_P):
                # Paragraf formatidagi savollar
                para = Paragraph(element, doc)
                if self._is_question_paragraph(para):
                    question = self._parse_paragraph_question(para, doc)
                    if question:
                        questions.append(question)

        # Agar jadvalda savol topilmasa, matnni boshqacha usulda parse qilish
        if not questions:
            questions = self._parse_text_questions(doc)

        return questions

    def _parse_table(self, table):
        """Jadval formatidagi savollarni parse qilish"""
        questions = []

        try:
            # Jadval qatorlarini o'qish
            for i, row in enumerate(table.rows):
                if i == 0:  # Header qatorini o'tkazib yuborish
                    continue

                cells = row.cells
                if len(cells) >= 5:
                    question = {
                        'id': cells[0].text.strip(),
                        'question': self._extract_cell_content(cells[1]),
                        'variants': {
                            'A': self._extract_cell_content(cells[2]),
                            'B': self._extract_cell_content(cells[3]),
                            'C': self._extract_cell_content(cells[4]),
                            'D': self._extract_cell_content(cells[5]) if len(cells) > 5 else {'text': '', 'images': []}
                        }
                    }
                    questions.append(question)

        except Exception as e:
            print(f"Jadval parse qilishda xatolik: {e}")

        return questions

    def _extract_cell_content(self, cell):
        """Katakdan matn va rasmlarni ajratib olish"""
        content = {
            'text': cell.text.strip(),
            'images': []
        }

        # Rasmlarni topish
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.element.xpath('.//pic:pic'):
                    # Rasm topildi
                    image_info = self._extract_image_from_run(run)
                    if image_info:
                        content['images'].append(image_info)

        return content

    def _extract_image_from_run(self, run):
        """Run dan rasm ma'lumotlarini ajratib olish"""
        try:
            # Rasm elementlarini topish
            blips = run.element.xpath('.//a:blip')
            if blips:
                rId = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId:
                    image_part = run.part.related_parts[rId]
                    image_data = image_part.blob

                    # Base64 ga aylantirish
                    image_base64 = base64.b64encode(image_data).decode('utf-8')

                    return {
                        'data': f"data:image/png;base64,{image_base64}",
                        'format': 'base64'
                    }
        except Exception as e:
            print(f"Rasm ajratishda xatolik: {e}")

        return None

    def _is_question_paragraph(self, para):
        """Paragraf savol ekanligini tekshirish"""
        text = para.text.strip()
        # Raqam bilan boshlanadigan paragraflarni savol sifatida qarash
        return bool(re.match(r'^\d+\.?\s+', text))

    def _parse_paragraph_question(self, para, doc):
        """Paragraf formatidagi savolni parse qilish"""
        text = para.text.strip()

        # Savol raqamini ajratish
        match = re.match(r'^(\d+)\.?\s+(.+)', text)
        if not match:
            return None

        question_id = match.group(1)
        question_text = match.group(2)

        # Variantlarni topish
        variants = self._find_variants_after_question(question_text)

        if variants:
            return {
                'id': question_id,
                'question': {'text': variants['question'], 'images': []},
                'variants': variants['options']
            }

        return None

    def _find_variants_after_question(self, text):
        """Savol matnidan variantlarni ajratish"""
        # A), B), C), D) formatidagi variantlarni topish
        variant_pattern = r'([A-D])\)\s*([^A-D]+?)(?=[A-D]\)|$)'
        matches = re.findall(variant_pattern, text, re.DOTALL)

        if len(matches) >= 2:
            options = {}
            question_text = text

            for letter, variant_text in matches:
                options[letter] = {
                    'text': variant_text.strip(),
                    'images': []
                }
                # Savoldan variantni olib tashlash
                question_text = re.sub(f'{letter}\\)\\s*{re.escape(variant_text.strip())}', '', question_text)

            return {
                'question': question_text.strip(),
                'options': options
            }

        return None

    def _parse_text_questions(self, doc):
        """Matn formatidagi savollarni parse qilish"""
        questions = []
        full_text = ""

        # Barcha matnni birlashtirish
        for para in doc.paragraphs:
            full_text += para.text + "\n"

        # Savollarni ajratish
        question_pattern = r'(\d+)\.?\s+(.+?)(?=\d+\.|$)'
        matches = re.findall(question_pattern, full_text, re.DOTALL)

        for question_id, question_text in matches:
            variants = self._find_variants_after_question(question_text.strip())

            if variants:
                question = {
                    'id': question_id,
                    'question': {'text': variants['question'], 'images': []},
                    'variants': variants['options']
                }
                questions.append(question)

        return questions

    def print_questions(self, questions_data):
        """Savollarni chiroyli formatda chop etish"""
        if not questions_data['success']:
            print(f"Xatolik: {questions_data['error']}")
            return

        questions = questions_data['questions']
        print(f"\nJami {len(questions)} ta savol topildi:\n")

        for q in questions:
            print(f"Savol {q['id']}: {q['question']['text']}")

            if q['question']['images']:
                print(f"  📷 {len(q['question']['images'])} ta rasm")

            for letter, variant in q['variants'].items():
                print(f"  {letter}) {variant['text']}")
                if variant['images']:
                    print(f"     📷 {len(variant['images'])} ta rasm")

            print("-" * 50)

    def save_to_json(self, questions_data, output_file):
        """Natijalarni JSON faylga saqlash"""
        import json

        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(questions_data, f, ensure_ascii=False, indent=2)
            print(f"Natijalar {output_file} faylga saqlandi")
        except Exception as e:
            print(f"Faylga saqlashda xatolik: {e}")


# Ishlatish namunasi
def main():
    # Word test faylini o'qish
    reader = WordTestReader()

    # Fayl manzilini kiriting
    file_path = input("Word fayl manzilini kiriting: ").strip()

    if not os.path.exists(file_path):
        print("Fayl topilmadi!")
        return

    # Faylni o'qish
    print("Fayl o'qilmoqda...")
    result = reader.read_test_file(file_path)

    # Natijalarni ko'rsatish
    reader.print_questions(result)

    # JSON faylga saqlash (ixtiyoriy)
    save_json = input("Natijalarni JSON faylga saqlashni xohlaysizmi? (y/n): ").lower()
    if save_json == 'y':
        output_file = input("JSON fayl nomini kiriting (masalan: test_questions.json): ").strip()
        if not output_file:
            output_file = "test_questions.json"
        reader.save_to_json(result, output_file)


if __name__ == "__main__":
    main()